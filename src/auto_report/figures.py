# -*- coding: utf-8 -*-

# Imports #####################################################################

from typing import Optional
import os
import pandas as pd
import numpy as np
import xarray as xr
from docx import Document
from docx.shared import Inches
import streamlit as st
import contextily as ctx
import rioxarray as rxr
import geopandas as gpd
from rashdf import RasPlanHdf
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import matplotlib.gridspec as gridspec
import matplotlib.ticker as ticker
from matplotlib.colors import ListedColormap, BoundaryNorm
from matplotlib.lines import Line2D
from mpl_toolkits.axes_grid1 import make_axes_locatable

import pygeohydro as gh
from pygeohydro import WBD
from shapely.geometry import Point
from pygeohydro import NWIS
from pynhd import HP3D, GeoConnex
import pygeoutils as geoutils
import warnings

from hy_river import (
    get_dem_data,
    get_nhd_flowlines,
    get_nwis,
    filter_nid,
)
from hdf_utils import get_model_perimeter, get_plan_cell_points
from metrics import calc_metrics
from tables import fill_calibration_metrics_table, fill_computation_settings_table

warnings.filterwarnings("ignore")

# assign a global font size for the plots
plt.rcParams.update({"font.size": 16})
# Set the display option for floating-point numbers to show only 3 decimal places
pd.options.display.float_format = "{:.3f}".format
# Functions ###################################################################


def calc_aep(data: pd.DataFrame):
    """
    Calculate the annual probability of exceedance

    Parameters
    ----------
    data : pd.DataFrame
        The data to calculate the probability of exceedance for

    Returns
    -------
    df_sorted : pd.DataFrame
        Columns: ['discharge', 'rank', 'pr']
        The data with a new column "pr" representing the percent probability
        (0-100) of exceedance for the given annual peak discharge.
    """
    # First ensure the data is formated correctly as a dataframe
    df = pd.DataFrame(data.values, index=data.index, columns=["discharge"])
    # Sort the values from smallest to largest
    df_sorted = df.sort_values(by="discharge")
    n = df_sorted.shape[0]
    # Assign a rank to each sorted value
    df_sorted.insert(0, "rank", range(1, 1 + n))
    # Determine the probability of exceedance
    df_sorted["pr"] = ((n - df_sorted["rank"] + 1) / (n + 1)) * 100
    # return the dataframe with a new column "pr" representing the probability of exceedance for the given value
    return df_sorted


def add_image_to_keyword(report_document: Document, keyword: str, image_path: str):
    """
    Add an image to a keyword within a document

    Parameters
    ----------
    report_document : Docx Document
        The document to modify
    keyword : str
        The keyword to search for in the document
    image_path : str
        The file path to the image to add

    Returns
    -------
    report_document
        The modified document

    """

    # Iterate through paragraphs
    for para in report_document.paragraphs:
        if keyword.lower() in para.text.lower():
            # Add image to a new paragraph before the keyword
            p = para.insert_paragraph_before()
            run = p.add_run()
            run.add_picture(image_path, width=Inches(6))

    # return the modified document
    return report_document


def plot_pilot_study_area(
    model_perimeter: gpd.GeoDataFrame,
    frequency_threshold: int,
    domain_name: str,
    root_dir: str,
    report_document: Optional[Document] = None,
    report_keywords: Optional[dict] = None,
):
    """
    Generate the Section 01 Figure 01 for the report
    Overview figure showing location of modeling unit relative to the HUC4 pilot project study area

    Parameters
    ----------
    model_perimeter : gpd.GeoDataFrame
        The perimeter of the model
    frequency_threshold : int
        The threshold frequency for the stream names. Ex: filter to streams whose
        name occurs 20 times or more within the NHDPlus HR network
    domain_name : str
        The name of the domain
    root_dir : str
        The root directory
    report_document : Docx Document
        The document to modify
    report_keywords : dict
        The keywords to search for in the document

    return
    ------
    if report_document is None and report_keywords is None:
        image_path : str
            The file path to the image
    else:
        report_document : Docx Document
            The modified document
        report_keywords : dict
            The updated values for the keywords
    """
    # determine the center of the model perimeter
    model_bbox = model_perimeter.bounds
    center = (
        model_bbox[["minx", "maxx"]].mean(axis=1).values[0],
        model_bbox[["miny", "maxy"]].mean(axis=1).values[0],
    )
    point = Point(center)

    ### HUC4 Boundary
    # Create an instance of the WBD class for HUC8
    wbd_huc4 = WBD("huc4")
    wbd_huc8 = WBD("huc8")
    # Query the HUC8 watershed boundary that covers the specified coordinates
    huc4_boundary = wbd_huc4.bygeom(point)
    huc8_boundary = wbd_huc8.bygeom(point)

    ### NHDPlus HR Network within the model perimeter
    # Create an instance of the NHDPlusHR class
    nhd3d = HP3D("flowline")
    # Query the NHDPlus HR network that covers the model perimeter
    geom_query = model_perimeter.union_all()
    network = nhd3d.bygeom(geom_query, model_perimeter.crs)
    # Filter the network to only include the mainstem waterbody connectors
    stream_names = network[network.featuretypelabel == "Waterbody Connector"][
        "gnisidlabel"
    ]
    # Determine the frequency of each stream name
    stream_name_freq = stream_names.value_counts()
    # Drop streams with a frequency less than the threshold
    stream_name_freq = stream_name_freq[stream_name_freq >= frequency_threshold]
    streams_df = network[network.featuretypelabel == "Waterbody Connector"]
    streams_df = streams_df[streams_df["gnisidlabel"].isin(stream_name_freq.index)]

    ### Mainstem Reach
    gcx = GeoConnex("mainstems")
    mainstem_reach_name = stream_name_freq.idxmax()
    # Determine the mainstem reach ID from the streams_df. The ID is the last part of the mainstemid url
    # Ex: https://geoconnex.us/ref/mainstems/322043
    mainstem_reach_id = (
        streams_df[streams_df["gnisidlabel"] == mainstem_reach_name]["mainstemid"]
        .unique()[0]
        .split("/")[-1]
    )
    mainstem_reach = gcx.byid("id", mainstem_reach_id)

    ### Generate the figure
    fig, ax = plt.subplots(figsize=(10, 10))
    # Add the model perimeter, HUC4 boundary, streams, and mainstem reach to the plot
    model_perimeter.plot(
        ax=ax, facecolor="red", edgecolor="black", alpha=0.2, linewidth=1
    )
    model_perimeter.plot(
        ax=ax, facecolor="none", edgecolor="black", alpha=1, linewidth=1, linestyle="--"
    )
    huc4_boundary.plot(ax=ax, facecolor="none", edgecolor="black", linewidth=3)
    streams_df.plot(ax=ax, color="blue", linewidth=1, alpha=0.5)
    mainstem_reach.plot(ax=ax, color="blue", linewidth=2, label=mainstem_reach_name)
    # Add lat/lon labels
    ax.set_xlabel("Longitude")
    ax.set_ylabel("Latitude")
    # Remove empty space around the plot
    plt.tight_layout()
    plt.margins(0)
    # Create custom legend handles
    custom_handles = [
        mpatches.Patch(
            facecolor="red", edgecolor="black", alpha=0.2, linestyle="--"
        ),  # Model Perimeter as a box
        Line2D([0], [0], color="black", lw=3),  # HUC4 Boundary
        Line2D([0], [0], color="blue", lw=3),  # Mainstem Reach
    ]
    labels = ["Model Domain", "HUC4 Boundary", mainstem_reach_name]

    # Create a divider for the existing axes instance
    divider = make_axes_locatable(ax)
    # Append axes to the bottom of ax, with 5% width of ax
    legend_ax = divider.append_axes("bottom", size="5%", pad=1)
    # Add the legend to the new axis
    legend_ax.axis("off")
    legend_ax.legend(custom_handles, labels, loc="center", framealpha=1, ncols=3)

    # Add a basemap
    ctx.add_basemap(
        ax=ax, crs=model_perimeter.crs, source=ctx.providers.OpenStreetMap.Mapnik
    )

    # Save the figure
    image_path = os.path.join(root_dir, f"{domain_name}_figure_pilot_study_area.png")
    fig.savefig(image_path, bbox_inches="tight")
    plt.close(fig)
    if report_document is None and report_keywords is None:
        return image_path
    else:
        # Search for the figure keyword within the document and add the image above it
        report_document = add_image_to_keyword(
            report_document, "«figure_pilot_study_area»", image_path
        )
        os.remove(image_path)
        # Update the report text
        basin_name, pilot_name = (
            huc8_boundary.name.values[0],
            huc4_boundary.name.values[0],
        )
        report_keywords["Model_Unit_Name"] = basin_name
        report_keywords["Header_Main"] = f"{pilot_name} / {basin_name}"
        report_keywords["Header_01"] = f"{pilot_name} / {basin_name}"
        report_keywords[
            "figure_pilot_study_area"
        ] = f"{pilot_name} Basin Overview with {basin_name} Modeling Unit"
        return report_document, report_keywords


def plot_dem(
    model_perimeter: gpd.GeoDataFrame,
    domain_name: str,
    root_dir: str,
    report_document: Optional[Document] = None,
    report_keywords: Optional[dict] = None,
):
    """
    Generate the Section 01 Figure 02 for the report
    Model perimeter with digital elevation model for terrain

    Parameters
    ----------
    model_perimeter : gpd.GeoDataFrame
        The perimeter of the model
    domain_name : str
        The name of the domain
    root_dir : str
        The root directory
    report_document : Docx Document
        The document to modify
    report_keywords : dict
        The keywords to search for in the document
    return
    ------
    if report_document is None and report_keywords is None:
        image_path : str
            The file path to the image
    else:
        report_document : Docx Document
            The modified document
        report_keywords : dict
            The updated values for the keywords
    """
    # Get the DEM data within the model perimeter
    dem = get_dem_data(model_perimeter)
    # Generate the figure
    fig, ax = plt.subplots(figsize=(10, 10), dpi=300)
    if dem is None:
        # Plot the model perimeter
        model_perimeter.plot(ax=ax, edgecolor="black", linewidth=3, facecolor="none")
        # set an x axis label
        ax.set_xlabel("Longitude")
        # set a y axis label
        ax.set_ylabel("Latitude")
        # remove the title
        ax.set_title("")
        # Remove empty space around the plot
        plt.tight_layout()
        # Add a basemap
        ctx.add_basemap(
            ax, crs=model_perimeter.crs, source=ctx.providers.OpenStreetMap.Mapnik
        )
        # Save the figure
        image_path = os.path.join(root_dir, f"{domain_name}_figure_dem.png")
        fig.savefig(image_path, bbox_inches="tight")
        plt.close(fig)
        if report_document is None and report_keywords is None:
            return image_path
        else:
            # Search for the keyword within the document and add the image above it
            report_document = add_image_to_keyword(
                report_document, "«figure_dem»", image_path
            )
            os.remove(image_path)
            # Update the report text
            basin_name = report_keywords["Model_Unit_Name"]
            report_keywords[
                "figure_dem"
            ] = f"{basin_name} Digital Elevation Model (DEM)"
            return report_document, report_keywords
    else:
        # Convert units from m to ft
        dem = dem * 3.28084
        # Plot the DEM
        cax = dem.plot(ax=ax, cmap="terrain", add_colorbar=False)
        model_perimeter.plot(ax=ax, edgecolor="black", linewidth=3, facecolor="none")
        # set an x and y axis labels
        ax.set_xlabel("Longitude")
        ax.set_ylabel("Latitude")
        # remove the title
        ax.set_title("")
        # Remove empty space around the plot
        plt.tight_layout()
        # Create a divider for the existing axes instance
        divider = make_axes_locatable(ax)
        # Append axes to the bottom of ax, with 5% width of ax
        cbar_ax = divider.append_axes("bottom", size="5%", pad=0.8)
        # Add the colorbar to the new axis
        cbar = fig.colorbar(cax, cax=cbar_ax, orientation="horizontal")
        # Set the colorbar label
        cbar.set_label("Elevation (ft)")
        # Save the figure
        image_path = os.path.join(root_dir, f"{domain_name}_figure_dem.png")
        fig.savefig(image_path, bbox_inches="tight")
        plt.close(fig)
        if report_document is None and report_keywords is None:
            return image_path
        else:
            # Search for the keyword within the document and add the image above it
            report_document = add_image_to_keyword(
                report_document, "«figure_dem»", image_path
            )
            os.remove(image_path)
            # Update the report text
            basin_name = report_keywords["Model_Unit_Name"]
            report_keywords[
                "figure_dem"
            ] = f"{basin_name} Digital Elevation Model (DEM)"
            return report_document, report_keywords


def plot_stream_network(
    model_perimeter: gpd.GeoDataFrame,
    df_gages_usgs: gpd.GeoDataFrame,
    domain_name: str,
    nid_parquet_file_path: str,
    nid_dam_height: int,
    root_dir: str,
    active_streamlit: bool,
    report_document: Optional[Document] = None,
    report_keywords: Optional[dict] = None,
):
    """
    Generate the Section 04 Figure 03 for the report
    Model perimeter with USGS gages and stream network

    Parameters
    ----------
    model_perimeter : gpd.GeoDataFrame
        The perimeter of the model
    df_gages_usgs : gpd.GeoDataFrame
        The USGS gages located within the model perimeter
    domain_name : str
        The name of the domain
    nid_parquet_file_path: str
        The file path to the parquet file containing the backup NID dam data
    nid_dam_height: int
        The height threshold for the NID dams
    root_dir : str
        The root directory
    active_streamlit : bool
        If true, print statements will be displayed in the Streamlit app
        If false, print statements will be displayed in the terminal
    report_document : Docx Document
        The document to modify
    report_keywords : dict
        The keywords to search for in the document

    return
    ------
    if report_document is None and report_keywords is None:
        image_path : str
            The file path to the image
    else:
        report_document : Docx Document
            The modified document
        report_keywords : dict
            The updated values for the keywords
    """
    # Generate the figure
    fig, ax = plt.subplots(figsize=(10, 10))
    model_perimeter.boundary.plot(ax=ax, color="black", linewidth=3, zorder=4)
    # Query all NHD streams within the domain
    streams = get_nhd_flowlines(model_perimeter)
    if isinstance(streams, str):
        num_streams = streams
        if active_streamlit:
            st.write(f"Error retrieving NHD streams. {streams}")
        else:
            print(f"Error retrieving NHD streams. {streams}")
    else:
        print("Successfully retrieved NHD streams")
        streams.plot(ax=ax, color="blue", linewidth=1, zorder=1)
        num_streams = len(streams["gnis_name"].unique())

    # Query all NID dams within the domain
    dams = filter_nid(nid_parquet_file_path, model_perimeter, nid_dam_height)
    if dams is not None:
        num_dams = len(dams)
        dams.plot(
            ax=ax, color="purple", edgecolor="k", markersize=100, marker="^", zorder=2
        )
    else:
        print(
            f"No dams found in the NID inventory with a height greater than {nid_dam_height} ft"
        )
        num_dams = 0

    # Check if the USGS gages are available to plot
    if len(df_gages_usgs) > 0:
        # add the gage locations to the plot
        df_gages_usgs.plot(ax=ax, color="red", markersize=100, marker="o", zorder=3)
        num_gages = len(df_gages_usgs)
    else:
        num_gages = "Data Unavailable"
    # set an x and y axis labels
    ax.set_xlabel("Longitude")
    ax.set_ylabel("Latitude")
    # Remove empty space around the plot
    plt.tight_layout()
    plt.margins(0)

    # Create custom legend handles
    custom_handles = [
        Line2D([0], [0], color="black", lw=3),  # Model Perimeter
        Line2D([0], [0], color="blue", lw=1),  # NHD Streams
        Line2D([0], [0], color="purple", lw=0, marker="^", markersize=10),  # NID Dams
        Line2D([0], [0], color="red", lw=0, marker="o", markersize=10),  # USGS Gages
    ]
    labels = [
        "Model Domain",
        f"NHD Streams: {num_streams}",
        f"NID Dams >= {nid_dam_height} ft: {num_dams}",
        f"USGS Gages: {num_gages}",
    ]

    # Create a divider for the existing axes instance
    divider = make_axes_locatable(ax)
    # Append axes to the bottom of ax, with 5% width of ax
    legend_ax = divider.append_axes("bottom", size="5%", pad=1)
    # Add the legend to the new axis
    legend_ax.axis("off")
    legend_ax.legend(custom_handles, labels, loc="center", framealpha=1, ncols=2)
    # Add a basemap
    ctx.add_basemap(
        ax, crs=model_perimeter.crs, source=ctx.providers.OpenStreetMap.Mapnik
    )
    # Save the figure
    image_path = os.path.join(root_dir, f"{domain_name}_figure_basin_datasets.png")
    fig.savefig(image_path, bbox_inches="tight")
    plt.close(fig)
    if report_document is None and report_keywords is None:
        return image_path
    else:
        # Search for the keyword within the document and add the image above it
        report_document = add_image_to_keyword(
            report_document, "«figure_basin_datasets»", image_path
        )
        os.remove(image_path)
        # Update the report text
        report_keywords[
            "figure_basin_datasets"
        ] = f"Model Perimeter with NHD Streams, NID Dams, and USGS Gages"
        return report_document, report_keywords


def plot_gage_por(
    df_gages_usgs: gpd.GeoDataFrame,
    domain_name: str,
    root_dir: str,
    report_document: Optional[Document],
    report_keywords: Optional[dict],
):
    """
    Generate the Section 04 Figure 04 for the report
    Gage each gage period of record (POR) streamflow data analytic summary

    Parameters
    ----------
    df_gages_usgs : gpd.GeoDataFrame
        The USGS gages located within the model perimeter
    domain_name : str
        The name of the domain
    root_dir : str
        The root directory
    report_document : Docx Document
        The document to modify
    report_keywords : dict
        The keywords to search for in the document

    return
    ------
    report_document : Docx Document
        The modified document
    report_keywords : dict
        The updated values for the keywords
    """
    if len(df_gages_usgs) > 0:
        # Create an instance of the NWIS class
        nwis = NWIS()
        generated_image_paths = []
        # Loop through each station
        for idx, row in df_gages_usgs.iterrows():
            station_id = row["site_no"]  # 08059590
            station = f"USGS-{station_id}"  # USGS-08059590
            station_name = row["station_nm"]  # Willow Creek at Highway 80
            print(f"Processing station {station} {station_name}")
            begin_date = row["begin_date"]
            end_date = row["end_date"]
            dates = (begin_date, end_date)
            # Get all available streamflow data within the specified date range
            qpor_df_daily = nwis.get_streamflow(
                [station_id], dates, mmd=False, freq="dv"
            )
            qpor_df_daily = qpor_df_daily * 35.3147  # Convert from cms to cfs

            # Calculate the monthly and annual peak streamflow data
            qpor_df_monthly = qpor_df_daily.copy()
            qpor_df_monthly = qpor_df_monthly.groupby(
                qpor_df_monthly.index.month
            ).mean()
            qpor_df_annual = qpor_df_daily.copy()
            qpor_df_annual = qpor_df_annual.groupby(qpor_df_annual.index.year).max()

            # Calculate the Annual Exceedance Probability for the annual peak streamflow data
            aep_df = calc_aep(qpor_df_annual)

            # Generate the figure
            fig = plt.figure(figsize=(15, 10))
            # Define a GridSpec with 2 rows and 3 columns
            gs = gridspec.GridSpec(2, 3, height_ratios=[1, 1])
            # Create subplots using the GridSpec layout
            ax1 = fig.add_subplot(gs[0, :])  # Top row, spans all columns
            ax2 = fig.add_subplot(
                gs[1, :2]
            )  # Bottom row, left subplot, spans two columns
            ax3 = fig.add_subplot(gs[1, 2])  # Bottom row, right subplot, single column

            # Plot the daily streamflow data
            ax1.plot(
                qpor_df_daily.index,
                qpor_df_daily.values,
                label="Daily Streamflow",
                color="blue",
                alpha=0.7,
            )
            # Plot the monthly average streamflow data
            qpor_df_monthly.plot(
                kind="bar",
                ax=ax2,
                color="blue",
                edgecolor="black",
                alpha=0.7,
                legend=False,
            )
            # Plot the annual peak streamflow data
            ax3.plot(
                aep_df.discharge.values,
                aep_df.pr.values,
                label="Annual Exceedance Probability",
                color="blue",
                linewidth=0,
                marker="o",
                alpha=0.7,
            )
            ax3.set_yscale("log")

            # Format the x-axis of ax2 to show month names titled 45 degrees
            ax2.set_xticklabels(
                [
                    "Jan",
                    "Feb",
                    "Mar",
                    "Apr",
                    "May",
                    "Jun",
                    "Jul",
                    "Aug",
                    "Sep",
                    "Oct",
                    "Nov",
                    "Dec",
                ],
                rotation=45,
            )
            # Format the x-axis of ax3 to tilt years 45 degrees
            ax3.set_xticklabels(np.round(aep_df.discharge.values, 1), rotation=45)

            # Set the custom y-axis formatter
            ax1.get_yaxis().set_major_formatter(
                ticker.FuncFormatter(lambda x, p: format(int(x), ","))
            )
            ax2.get_yaxis().set_major_formatter(
                ticker.FuncFormatter(lambda x, p: format(int(x), ","))
            )
            ax3.get_xaxis().set_major_formatter(
                ticker.FuncFormatter(lambda x, p: format(int(x), ","))
            )

            # Add a letter to each subplot
            ax1.text(0.01, 0.90, "(a)", transform=ax1.transAxes, fontsize=24)
            ax2.text(0.01, 0.90, "(b)", transform=ax2.transAxes, fontsize=24)
            ax3.text(0.85, 0.90, "(c)", transform=ax3.transAxes, fontsize=24)

            # Add gridlines to each subplot
            ax1.grid(True)
            ax2.grid(True)
            ax3.grid(True, which="both")

            # Add y-axis labels to each subplot
            ax1.set_ylabel("Q Daily (cfs)")
            ax2.set_ylabel("Q Monthly (cfs)")
            ax3.set_ylabel("AEP (%)")
            ax3.set_xlabel("Q Annual (cfs)")

            # Add a title to the figure
            fig.suptitle(f"{station} {station_name}", fontsize=24)
            # Display the plot
            plt.tight_layout()

            # Save the figure
            image_path = os.path.join(
                root_dir, f"{domain_name}_figure_stream_gage_summary_{station}.png"
            )
            fig.savefig(image_path, bbox_inches="tight")
            plt.close(fig)
            if report_document is None and report_keywords is None:
                generated_image_paths.append(image_path)
            else:
                # Search for the keyword within the document and add the image above it
                report_document = add_image_to_keyword(
                    report_document, "«figure_stream_gage_summary»", image_path
                )
                os.remove(image_path)

        if report_document is None and report_keywords is None:
            return generated_image_paths
        else:
            # Update the report text
            report_keywords[
                "figure_stream_gage_summary"
            ] = "USGS Gage Streamflow Summary Analytics"
            return report_document, report_keywords
    else:
        print("USGS gages are currently unavailable")
        if report_document is None and report_keywords is None:
            return None
        else:
            return report_document, report_keywords


def plot_nlcd(
    hdf_geom_file_path: str,
    nlcd_file_path: str,
    domain_name: str,
    root_dir: str,
    report_document: Optional[Document] = None,
    report_keywords: Optional[dict] = None,
):
    """
    Generate the Section 04 Figure 05 for the report
    Model perimeter with land cover usage from the NLCD dataset

    Parameters
    ----------
    hdf_geom_file_path : str
        The file path to the HDF5 geometry file
    nlcd_file_path : str
        The file path to the NLCD dataset
    domain_name : str
        The name of the domain
    root_dir : str
        The root directory
    report_document : Docx Document
        The document to modify
    report_keywords : dict
        The keywords to search for in the document

    return
    ------
    if report_document is None and report_keywords is None:
        image_path : str
            The file path to the image
    else:
        report_document : Docx Document
            The modified document
        report_keywords : dict
            The updated values for the keywords
    """
    # Create a dictionary of the land cover metadata
    meta = {
        11: "Open Water",
        12: "Perennial Ice/Snow",
        21: "Developed, Open Space",
        22: "Developed, Low Intensity",
        23: "Developed, Medium Intensity",
        24: "Developed, High Intensity",
        31: "Barren Land (Rock/Sand/Clay)",
        41: "Deciduous Forest",
        42: "Evergreen Forest",
        43: "Mixed Forest",
        51: "Dwarf Scrub",
        52: "Shrub/Scrub",
        71: "Grassland/Herbaceous",
        72: "Sedge/Herbaceous",
        73: "Lichens",
        74: "Moss",
        81: "Pasture/Hay",
        82: "Cultivated Crops",
        90: "Woody Wetlands",
        95: "Emergent Herbaceous Wetlands",
    }
    # Define the color pallet for the NLCD classes
    colors = {
        11: "#486DA2",  # Open Water
        12: "#E7EFFC",  # Perennial Ice/Snow
        21: "#E1CDCE",  # Developed, Open Space
        22: "#DC9881",  # Developed, Low Intensity
        23: "#F10100",  # Developed, Medium Intensity
        24: "#AB0101",  # Developed, High Intensity
        31: "#B3AFA4",  # Barren Land (Rock/Sand/Clay)
        41: "#6CA966",  # Deciduous Forest
        42: "#1D6533",  # Evergreen Forest
        43: "#BDCC93",  # Mixed Forest
        51: "#B49E48",  # Dwarf Scrub
        52: "#D1BB82",  # Shrub/Scrub
        71: "#EDECCD",  # Grassland/Herbaceous
        72: "#D0D181",  # Sedge/Herbaceous
        73: "#A4CC51",  # Lichens
        74: "#82BA9D",  # Moss
        81: "#DDD83E",  # Pasture/Hay
        82: "#AE7229",  # Cultivated Crops
        90: "#BBD7ED",  # Woody Wetlands
        95: "#71A4C1",  # Emergent Herbaceous Wetlands
    }

    # Load the file into an xarray Dataset object
    nlcd = rxr.open_rasterio(nlcd_file_path)
    nlcd = xr.Dataset({"data": nlcd})

    # Load the model perimeter from the HDF5 file
    model_perimeter = get_model_perimeter(
        hdf_geom_file_path, domain_name, project_to_4326=False
    )
    if model_perimeter.crs != nlcd.rio.crs:
        model_perimeter = model_perimeter.to_crs(nlcd.rio.crs)

    # Filter the colors and classes to only what is present within the provided dataset
    unique_nlcd = np.unique(nlcd.data.values)
    # seperate all unique values that do not belong to the NLCD classes
    non_nlcd = [value for value in unique_nlcd if value not in meta.keys()]
    # remove the non-NLCD values from the unique NLCD values
    unique_nlcd = [value for value in unique_nlcd if value not in non_nlcd]
    if len(unique_nlcd) == 0:
        raise ValueError("Provided NLCD raster does not contain any NLCD classes")

    colors = [colors[key] for key in unique_nlcd]
    nlcd_classes = pd.Series(meta)
    nlcd_classes = nlcd_classes.loc[unique_nlcd]

    # mask non-NLCD values from the dataset
    nlcd = nlcd.where(nlcd.isin(unique_nlcd))

    # Create a ListedColormap and BoundaryNorm
    cmap = ListedColormap(colors)
    norm = BoundaryNorm(list(nlcd_classes.keys()) + [100], cmap.N)
    levels = nlcd_classes.index

    fig, ax = plt.subplots(figsize=(10, 8), dpi=300)
    model_perimeter.plot(
        ax=ax, edgecolor="black", facecolor="none", linewidth=3, zorder=1
    )
    nlcd.data.plot(ax=ax, add_colorbar=False, cmap=cmap, norm=norm, zorder=0)
    # Create custom legend handles with edge color
    legend_handles = [
        mpatches.Patch(
            facecolor=cmap(norm(level)),
            label=f"{nlcd_classes[level]}",
            edgecolor="black",
        )
        for level in levels
    ]
    ax.set_title("")
    # Move the legend outside the plot
    ax.legend(
        loc="upper left",
        bbox_to_anchor=(1, 1),
        handles=legend_handles,
        title=f"NLCD Land Cover Usage",
    )
    # remove the axis labels
    ax.set_axis_off()
    # remove the axis ticks
    ax.set_xticks([])
    ax.set_yticks([])

    # Save the figure
    image_path = os.path.join(root_dir, f"{domain_name}_figure_nlcd.png")
    fig.savefig(image_path, bbox_inches="tight")
    plt.close(fig)
    if report_document is None and report_keywords is None:
        return image_path
    else:
        # Search for the keyword within the document and add the image above it
        report_document = add_image_to_keyword(
            report_document, "«figure_nlcd»", image_path
        )
        os.remove(image_path)
        # Update the report text
        report_keywords["figure_nlcd"] = f"NLCD Land Cover Usage"
        return report_document, report_keywords


def plot_soil_porosity(
    report_document: Document,
    report_keywords: dict,
    model_perimeter: gpd.GeoDataFrame,
    domain_name: str,
    root_dir: str,
):
    """
    Generate the Section 04 Figure 06 for the report
    Plot the soil porosity (inches/ft) within 1 ft of the soil profile.

    Parameters
    ----------
    report_document : Docx Document
        The document to modify
    report_keywords : dict
        The keywords to search for in the document
    model_perimeter : gpd.GeoDataFrame
        The perimeter of the model
    domain_name : str
        The name of the domain
    root_dir : str
        The root directory

    return
    ------
    report_document : Docx Document
        The modified document
    report_keywords : dict
        The updated values for the keywords
    """

    # Acquire the soil porosity dataset for the model perimeter
    por = gh.soil_properties("por")
    por = geoutils.xarray_geomask(
        por, model_perimeter.geometry.iloc[0], model_perimeter.crs
    )
    por = por.where(por.porosity > por.porosity.rio.nodata)
    por["por"] = por.porosity.rio.write_nodata(np.nan)

    # Generate the figure
    fig, ax = plt.subplots(figsize=(10, 10), dpi=300)
    # Plot the
    cax = por.por.plot(ax=ax, cmap="gist_earth_r", add_colorbar=False)
    model_perimeter.plot(ax=ax, edgecolor="black", linewidth=3, facecolor="none")
    # set an x and y axis labels
    ax.set_xlabel("Longitude")
    ax.set_ylabel("Latitude")
    # remove the title
    ax.set_title("")
    # Remove empty space around the plot
    plt.tight_layout()
    # Create a divider for the existing axes instance
    divider = make_axes_locatable(ax)
    # Append axes to the right of ax, with 5% width of ax
    cbar_ax = divider.append_axes("bottom", size="5%", pad=0.8)
    # Add the colorbar to the new axis
    cbar = fig.colorbar(cax, cax=cbar_ax, orientation="horizontal")
    # Set the colorbar label
    cbar.set_label("Soil Porosity (mm/m)")
    # Add a basemap
    ctx.add_basemap(
        ax, crs=model_perimeter.crs, source=ctx.providers.OpenStreetMap.Mapnik
    )
    # Save the figure
    image_path = os.path.join(root_dir, f"{domain_name}_figure_soils.png")
    fig.savefig(image_path, bbox_inches="tight")
    plt.close(fig)
    # Search for the keyword within the document and add the image above it
    report_document = add_image_to_keyword(
        report_document, "«figure_soils»", image_path
    )
    os.remove(image_path)
    # Update the report text
    report_keywords[
        "figure_soils"
    ] = "Soil Porosity (inches/ft) within 1 ft of the soil profile"
    return report_document, report_keywords


def plot_model_mesh(
    model_perimeter: gpd.GeoDataFrame,
    model_breaklines: gpd.GeoDataFrame,
    model_cells: gpd.GeoDataFrame,
    domain_name: str,
    root_dir: str,
    report_document: Optional[Document] = None,
    report_keywords: Optional[dict] = None,
):
    """
    Generate the Section 04 Figure 07 for the report
    Plot the final constructed mesh geometry of the model

    Parameters
    ----------
    model_perimeter : gpd.GeoDataFrame
        The perimeter of the model
    model_breaklines : gpd.GeoDataFrame
        The breaklines of the model
    model_cells : gpd.GeoDataFrame
        The 2D cells of the model
    domain_name : str
        The name of the domain
    root_dir : str
        The root directory
    report_document : Docx Document
        The document to modify
    report_keywords : dict
        The keywords to search for in the document

    return
    ------
    if report_document is None and report_keywords is None:
        image_path : str
            The file path to the image
    else:
        report_document : Docx Document
            The modified document
        report_keywords : dict
            The updated values for the keywords
    """
    num_cells = len(model_cells)
    num_breaklines = len(model_breaklines)
    # create a plot of the model geometry
    fig, ax = plt.subplots(figsize=(10, 10), dpi=300)
    model_perimeter.plot(ax=ax, facecolor="none", edgecolor="black", linewidth=3)
    if num_cells > 0:
        model_cells.plot(ax=ax, facecolor="none", edgecolor="blue", linewidth=0.25)
    if num_breaklines > 0:
        model_breaklines.plot(ax=ax, color="red", linewidth=1, alpha=0.5)

    # Create custom legend handles
    custom_handles = [
        Line2D([0], [0], color="black", lw=3),  # Model Domain
        Line2D([0], [0], color="red", lw=1, alpha=0.5),  # Breaklines
        mpatches.Patch(facecolor="none", edgecolor="blue"),  # 2D Cells
    ]
    labels = [
        "Model Domain",
        f"Breaklines: {num_breaklines:,}",
        f"2D Cells: {num_cells:,}",
    ]

    # Create a divider for the existing axes instance
    divider = make_axes_locatable(ax)
    # Append axes to the bottom of ax, with 5% width of ax
    legend_ax = divider.append_axes("bottom", size="5%", pad=1)
    # Add the legend to the new axis
    legend_ax.axis("off")
    legend_ax.legend(custom_handles, labels, loc="center", framealpha=1, ncols=3)

    # Add a basemap
    ctx.add_basemap(
        ax, crs=model_perimeter.crs, source=ctx.providers.OpenStreetMap.Mapnik
    )
    ax.set_xlabel("Longitude")
    ax.set_ylabel("Latitude")
    # Save the figure
    image_path = os.path.join(root_dir, f"{domain_name}_figure_model_mesh.png")
    fig.savefig(image_path, bbox_inches="tight")
    plt.close(fig)
    if report_document is None and report_keywords is None:
        return image_path
    else:
        # Search for the keyword within the document and add the image above it
        report_document = add_image_to_keyword(
            report_document, "«figure_model_mesh»", image_path
        )
        os.remove(image_path)
        # Update the report text
        report_keywords[
            "figure_model_mesh"
        ] = "Final Constructed Mesh Geometry of the Model"
        return report_document, report_keywords


def calc_wse_stats(df: pd.DataFrame, target_column: str):
    """ "
    Calculate the frequency, PDF, and CDF of a target column in a dataframe

    Parameters
    ----------
    df : pd.DataFrame
        A pandas dataframe with the target column
    target_column : str
        The name of the target column in the dataframe

    Returns
    -------
    stats_df : pd.DataFrame
        A pandas dataframe with the frequency, PDF, and CDF of the target column
    """
    # filter out values of -9999 and zeros
    df = df[df[target_column] != -9999]
    df = df[df[target_column] > 0]
    # Frequency
    stats_df = (
        df.groupby(target_column)[target_column]
        .agg("count")
        .pipe(pd.DataFrame)
        .rename(columns={target_column: "frequency"})
    )
    # PDF
    stats_df["pdf"] = stats_df["frequency"] / sum(stats_df["frequency"])
    # CDF
    stats_df["cdf"] = stats_df["pdf"].cumsum() * 100
    stats_df = stats_df.reset_index()
    return stats_df


def plot_hist_cdf(
    stats_df: pd.DataFrame,
    target_column: str,
    plot_title: str,
    threshold: float,
    num_bins: int,
    output_path: str,
):
    """
    Plot a histogram and CDF of a target column in a dataframe

    Parameters
    ----------
    stats_df : pd.DataFrame
        A pandas dataframe with the frequency, PDF, and CDF of the target column
    target_column : str
        The name of the target column in the dataframe
    plot_title : str
        The title of the plot
    threshold : float
        The threshold for the histogram's x-axis
    num_bins : int
        The number of bins for the histogram
    output_path : str
        The path to save the plot
    """

    fig, ax = plt.subplots(2, 1, sharex=True, figsize=(10, 10))
    ax = ax.flatten()
    # plot the histogram
    target_values = stats_df[target_column].values
    filtered_values = target_values[(target_values <= threshold) & (target_values >= 0)]
    weights = np.ones_like(filtered_values) / float(len(filtered_values)) * 100
    n, bins, patches = ax[0].hist(
        filtered_values,
        bins=num_bins,
        weights=weights,
        range=(0, threshold),
        density=False,
    )
    # determine the percentage of cells that have errors greater than the threshold
    if target_column == "ttp_hrs":
        # cells with time to peak occuring within the last hour of the simulation
        exceedance = int(len(target_values[target_values >= threshold]))
        plot_units = "hours"
    else:
        # cells with WSE errors greater than the threshold
        exceedance = int(len(target_values[(target_values > threshold)]))
        plot_units = "ft"

    ax[0].set(
        xlabel=f"{plot_title} ({plot_units})",
        ylabel="Cell Count (%)",
        title=f"Histogram: n bins = {num_bins}",
    )
    ax[0].set_xlim(0, threshold)
    ax[0].grid()
    # add a legend
    if target_column == "ttp_hrs":
        ax[0].legend(
            [f"{plot_title} >= {threshold:.0f}: {exceedance} Cells"], loc="upper right"
        )
    else:
        ax[0].legend(
            [f"{plot_title} > {threshold:.1f}: {exceedance} Cells"], loc="upper right"
        )

    # plot the CDF
    cdf_df = stats_df[
        (stats_df[target_column] <= threshold) & (stats_df[target_column] >= 0)
    ]
    # Resample the data to num_bins equally spaced points
    x = np.linspace(cdf_df[target_column].min(), cdf_df[target_column].max(), num_bins)
    y = np.interp(
        x, cdf_df[target_column].sort_values(), np.linspace(0, 100, len(cdf_df))
    )
    cdf_df = pd.DataFrame({target_column: x, "cdf": y})

    ax[1].plot(cdf_df[target_column], cdf_df["cdf"], linewidth=3)
    ax[1].set(
        xlabel=f"{plot_title} ({plot_units})",
        ylabel="P(X <= x) (%)",
        title="Non-Exceedance Probability",
    )
    ax[1].set_xlim(0, threshold)
    ax[1].set_ylim(0, 100)
    ax[1].grid()
    if target_column == "ttp_hrs":
        ax[1].xaxis.set_major_formatter(
            plt.FuncFormatter(lambda x, _: "{:.0f}".format(x))
        )
    else:
        ax[1].xaxis.set_major_formatter(
            plt.FuncFormatter(lambda x, _: "{:.2f}".format(x))
        )
    # add space between both subplots
    plt.tight_layout()
    # save the plot
    plt.savefig(output_path)
    plt.close(fig)


def plot_wse_errors(
    hdf_plan_file_path: str,
    wse_error_threshold: float,
    num_bins: int,
    root_dir: str,
    plan_index: Optional[int] = None,
    report_document: Optional[Document] = None,
    report_keywords: Optional[dict] = None,
):
    """
    Generate the Appendix A Figure 09 for the report
    Max WSE Error QC

    Parameters
    ----------
    hdf_plan_file_path : str
        The path to the HDF plan file
    wse_error_threshold: float
        The threshold for the WSE error
    num_bins: int
        The number of bins for the histogram
    root_dir : str
        The root directory
    plan_index: int
        The index of the plan to use
    report_document : Docx Document
        The document to modify
    report_keywords : dict
        The keywords to search for in the document

    Returns
    -------
    if report_document is None and report_keywords is None:
        image_path : str
            The file path to the image
    else:
        report_document : Docx Document
            The modified document
        report_keywords : dict
            The updated values for the keywords
    """
    # Get the mesh cell points and domain name
    cell_points = get_plan_cell_points(hdf_plan_file_path)
    domain_name = cell_points["mesh_name"].unique()[0]
    # Generate the Figure for the WSE Error QC
    image_path = os.path.join(root_dir, f"{domain_name}_figure_wse_errors.png")

    # Process Max WSE Errors and export graphics
    if len(cell_points) == 0:
        print("No cells found in the HDF file")
        output_message = "No cells found in the HDF file"
        if report_document is None and report_keywords is None and plan_index is None:
            return output_message
        else:
            report_keywords[
                f"plan0{plan_index}_figure_wse_errors"
            ] = "No cells found in the HDF file"
            return report_document, report_keywords
    elif len(cell_points[cell_points["max_ws_err"] > 0]) == 0:
        print("No cells with WSE errors greater than zero")
        output_message = "No cells with WSE errors greater than zero"
        if report_document is None and report_keywords is None and plan_index is None:
            return output_message
        else:
            report_keywords[
                f"plan0{plan_index}_figure_wse_errors"
            ] = "No cells with WSE errors greater than zero"
            return report_document, report_keywords
    else:
        # Calculate the frequency, PDF, and CDF of the max WSE errors
        cell_stats_df = calc_wse_stats(cell_points, "max_ws_err")
        num_cells = len(cell_points)
        cells_exceeding_threshold = cell_points[
            cell_points["max_ws_err"] > wse_error_threshold
        ]
        num_cells_exceeding_threshold = len(cells_exceeding_threshold)
        max_wse_error = cell_points["max_ws_err"].max()
        # Plot the histogram and CDF of the max WSE errors
        plot_hist_cdf(
            cell_stats_df,
            "max_ws_err",
            "Max WSE Error",
            wse_error_threshold,
            num_bins,
            image_path,
        )
        if report_document is None and report_keywords is None and plan_index is None:
            return image_path
        else:
            # Search for the keyword within the document and add the image above it
            report_document = add_image_to_keyword(
                report_document, f"«plan0{plan_index}_figure_wse_errors»", image_path
            )
            os.remove(image_path)
            # Update the report text
            report_keywords[
                f"plan0{plan_index}_figure_wse_errors"
            ] = f"""Maximum Water Surface Elevation Error Spatial Distribution. Among the {num_cells:,} wetted cells, approximately {num_cells_exceeding_threshold:,} cells have errors exceeding a threshold of {wse_error_threshold:.2f}-feet, with a maximum recorded error of {max_wse_error:.1f}-feet.
            """
            return report_document, report_keywords


def plot_wse_ttp(
    hdf_plan_file_path: str,
    num_bins: int,
    root_dir: str,
    plan_index: Optional[int] = None,
    report_document: Optional[Document] = None,
    report_keywords: Optional[dict] = None,
):
    """
    Generate the Appendix A Figure 10 for the report
    Time-to-Peak WSE QC

    Parameters
    ----------
    hdf_plan_file_path : str
        The path to the HDF plan file
    num_bins: int
        The number of bins for the histogram
    root_dir : str
        The root directory
    plan_index: int
        The index of the plan to use
    report_document : Docx Document
        The document to modify
    report_keywords : dict
        The keywords to search for in the document

    Returns
    -------
    if report_document is None and report_keywords is None:
        output_message : str
            The output message
    else:
        report_document : Docx Document
            The modified document
        report_keywords : dict
            The updated values for the keywords
    """
    # Get the mesh cell points and domain name
    cell_points = get_plan_cell_points(hdf_plan_file_path)
    domain_name = cell_points["mesh_name"].unique()[0]
    # Generate the Figure for the WSE Error QC
    image_path = os.path.join(root_dir, f"{domain_name}_figure_wse_ttp.png")

    # Process WSE time to peak
    if len(cell_points) == 0:
        print("No cells found in the HDF file")
        output_message = "No cells found in the HDF file"
        if report_document is None and report_keywords is None and plan_index is None:
            return output_message
        else:
            report_keywords[
                f"plan0{plan_index}_figure_wse_ttp"
            ] = "No cells found in the HDF file"
            return report_document, report_keywords

    # Determine the global min time to peak for the simulation start time
    start_time = cell_points["min_ws_time"].min()
    # find the time to peak between the max and min water surface elevation times
    ttp = cell_points["max_ws_time"] - start_time
    # convert to hours
    ttp = ttp.dt.total_seconds() / 3600
    ttp = ttp.to_frame()
    ttp.columns = ["ttp_hrs"]

    if len(ttp[ttp["ttp_hrs"] > 0]) == 0:
        print("No cells with time to peak greater than zero")
        output_message = "No cells with time to peak greater than zero"
        if report_document is None and report_keywords is None and plan_index is None:
            return output_message
        else:
            report_keywords[
                f"plan0{plan_index}_figure_wse_ttp"
            ] = "No cells with time to peak greater than zero"
            return report_document, report_keywords
    else:
        # Define the global max time to peak
        ttp_max = ttp["ttp_hrs"].max()
        num_cells = len(cell_points)
        cell_points = None
        # Calculate the frequency, PDF, and CDF of the time to peak
        cell_stats_df = calc_wse_stats(ttp, "ttp_hrs")
        cells_exceeding_threshold = cell_stats_df[cell_stats_df["ttp_hrs"] >= ttp_max]
        num_cells_exceeding_threshold = len(cells_exceeding_threshold)
        # Plot the histogram and CDF of the time to peak
        plot_hist_cdf(
            cell_stats_df,
            "ttp_hrs",
            "Time to Peak",
            ttp_max,
            num_bins,
            image_path,
        )
        if report_document is None and report_keywords is None and plan_index is None:
            return image_path
        else:
            # Search for the keyword within the document and add the image above it
            report_document = add_image_to_keyword(
                report_document, f"«plan0{plan_index}_figure_wse_ttp»", image_path
            )
            os.remove(image_path)
            # Update the report text
            report_keywords[
                f"plan0{plan_index}_figure_wse_ttp"
            ] = f"""Water Surface Elevation Time-to-Peak Spatial Distribution. The full duration for the model simulation is {ttp_max:.0f} hours. Among the {num_cells:,} wetted cells, at least {num_cells_exceeding_threshold:,} was still increasing in water surface elevation during the final hour of the simulation.
            """
            return report_document, report_keywords


def find_timstep_freq(df: pd.DataFrame):
    """
    Find the timestep frequency of the DataFrame

    Parameters
    ----------
    df : pd.DataFrame
        The DataFrame to analyze with a datetime index

    Returns
    -------
    ts_freq : pd.Timedelta
    """
    # Check if the index is a MultiIndex
    if isinstance(df.index, pd.MultiIndex):
        ts_index = df.index.droplevel(1)
        ts_freq = ts_index[1] - ts_index[0]
        # cast the frequency to a Timedelta object
        ts_freq = pd.Timedelta(ts_freq)
        return ts_freq
    else:
        ts_freq = df.index[1] - df.index[0]
        # cast the frequency to a Timedelta object
        ts_freq = pd.Timedelta(ts_freq)
        return ts_freq


def format_datetime(qobs_df: pd.DataFrame, qsim_df: pd.DataFrame):
    """
    Format the datetime index of the observed and modeled datasets
    and resample the data to the same timestep frequency.

    Parameters
    ----------
    qobs_df : pd.DataFrame
        The observed data DataFrame
    qsim_df : pd.DataFrame
        The modeled data DataFrame

    Returns
    -------
    qobs_df : pd.DataFrame
        The formatted observed data DataFrame
    qsim_df : pd.DataFrame
        The formatted modeled data DataFrame
    timestep : str
        The timestep frequency of the gage data
    """
    # Determine the frequency of the data
    qobs_freq = find_timstep_freq(qobs_df)
    qsim_freq = find_timstep_freq(qsim_df)

    # Resample the data to the same timestep frequency of the model
    if qobs_freq < qsim_freq:
        print(f"Resampling the observed data from {qobs_freq} to {qsim_freq}")
        qobs_df = qobs_df.resample(qsim_freq).mean()
        if qsim_freq.days == 1:
            timestep = f"{qsim_freq.days}-d"
            qobs_df.index = pd.to_datetime(qobs_df.index.date, format="%Y-%m-%d")
            qsim_df.index = pd.to_datetime(qsim_df.index.date, format="%Y-%m-%d")
            print(f"Data is at a {timestep} timestep")
            return qobs_df, qsim_df, timestep
        if qsim_freq.seconds/60 < 60:
            timestep = f"{qsim_freq.seconds // 60}-min"
            qobs_df.index = pd.to_datetime(qobs_df.index, format="%Y-%m-%d %H:%M")
            qsim_df.index = pd.to_datetime(qsim_df.index, format="%Y-%m-%d %H:%M")
            print(f"Data is at a {timestep} timestep")
            return qobs_df, qsim_df, timestep
        if qsim_freq.seconds/60 >= 60:
            timestep = f"{qsim_freq.seconds // 3600}-hr"
            qobs_df.index = pd.to_datetime(qobs_df.index, format="%Y-%m-%d %H")
            qsim_df.index = pd.to_datetime(qsim_df.index, format="%Y-%m-%d %H")
            print(f"Data is at a {timestep} timestep")
            return qobs_df, qsim_df, timestep

    elif qobs_freq > qsim_freq:
        print(f"Resampling the modeled data from {qsim_freq} to {qobs_freq}")
        qsim_df = qsim_df.resample(qobs_freq).mean()
        if qobs_freq.days == 1:
            timestep = f"{qobs_freq.days}-d"
            qobs_df.index = pd.to_datetime(qobs_df.index.date, format="%Y-%m-%d")
            qsim_df.index = pd.to_datetime(qsim_df.index.date, format="%Y-%m-%d")
            print(f"Data is at a {timestep} timestep")
            return qobs_df, qsim_df, timestep
        if qobs_freq.seconds/60 < 60 and qobs_freq.days != 1:
            timestep = f"{qobs_freq.seconds // 60}-min"
            qobs_df.index = pd.to_datetime(qobs_df.index, format="%Y-%m-%d %H:%M")
            qsim_df.index = pd.to_datetime(qsim_df.index, format="%Y-%m-%d %H:%M")
            print(f"Data is at a {timestep} timestep")
            return qobs_df, qsim_df, timestep
        if qobs_freq.seconds/60 >= 60 and qobs_freq.days != 1:
            timestep = f"{qobs_freq.seconds // 3600}-hr"
            qobs_df.index = pd.to_datetime(qobs_df.index, format="%Y-%m-%d %H")
            qsim_df.index = pd.to_datetime(qsim_df.index, format="%Y-%m-%d %H")
            print(f"Data is at a {timestep} timestep")
            return qobs_df, qsim_df, timestep

    else:
        print(
            f"Observed and modeled data are at the same timestep frequency of {qobs_freq}"
        )
        if qobs_freq.days == 1:
            timestep = f"{qobs_freq.days}-d"
            qobs_df.index = pd.to_datetime(qobs_df.index.date, format="%Y-%m-%d")
            qsim_df.index = pd.to_datetime(qsim_df.index.date, format="%Y-%m-%d")
            print(f"Data is at a {timestep} timestep")
            return qobs_df, qsim_df, timestep
        if qobs_freq.seconds/60 < 60:
            timestep = f"{qobs_freq.seconds // 60}-min"
            qobs_df.index = pd.to_datetime(qobs_df.index, format="%Y-%m-%d %H:%M")
            qsim_df.index = pd.to_datetime(qsim_df.index, format="%Y-%m-%d %H:%M")
            print(f"Data is at a {timestep} timestep")
            return qobs_df, qsim_df, timestep
        if qobs_freq.seconds/60 >= 60:
            timestep = f"{qobs_freq.seconds // 3600}-hr"
            qobs_df.index = pd.to_datetime(qobs_df.index, format="%Y-%m-%d %H")
            qsim_df.index = pd.to_datetime(qsim_df.index, format="%Y-%m-%d %H")
            print(f"Data is at a {timestep} timestep")
            return qobs_df, qsim_df, timestep
    


def plot_hydrographs(
    hdf_plan_file_path: str,
    df_gages_usgs: gpd.GeoDataFrame,
    parameter: str,
    domain_name: str,
    root_dir: str,
    plan_index: Optional[int] = None,
    report_document: Optional[Document] = None,
    report_keywords: Optional[dict] = None,
):
    """
    Generate the gage hydrograph calibration plots

    Parameters
    ----------
    hdf_plan_file_path : str
        The path to the HDF plan file
    df_gages_usgs : gpd.GeoDataFrame
        The USGS gages located within the model perimeter
    parameter : str
        The parameter to plot. One of 'Flow' or 'Stage'
    domain_name : str
        The name of the domain
    root_dir : str
        The root directory
    plan_index : int
        The index of the plan HDF file being used. One of [1, 2, 3, 4, 5, 6]
    report_document : Docx Document
        The document to modify
    report_keywords : dict
        The keywords to search for in the document

    Returns
    -------
    if report_document is None and report_keywords is None:
        image_path_list : list
            The list of file paths to the images
    else:
        report_document : Docx Document
            The modified document
        report_keywords : dict
            The updated values for the keywords
    """

    # Open the HDF plan file for the reference line data
    try:
        plan_hdf = RasPlanHdf.open_uri(hdf_plan_file_path)
    except Exception as e:
        raise FileNotFoundError(
            f"The provided HDF plan file {hdf_plan_file_path} does not exist. Please verify the file path."
        ) from e
    ref_lines = plan_hdf.reference_lines()
    ref_lines = ref_lines.to_crs(epsg=4326)  # convert to EPSG:4326
    ref_lines_ds = plan_hdf.reference_lines_timeseries_output()
    # Get the simulation start and end times
    plan_params = plan_hdf.get_plan_param_attrs()
    plan_attrs = plan_hdf.get_plan_info_attrs()
    start_time = plan_attrs["Simulation Start Time"]
    end_time = plan_attrs["Simulation End Time"]
    path_start_date = start_time.strftime("%Y-%m-%d")
    path_end_date = end_time.strftime("%Y-%m-%d")
    print(f"Calibration period: {path_start_date} to {path_end_date}")

    if parameter == "Flow":
        y_label_txt = "Streamflow (cfs)"
        sim_parameter = "Flow"
        obs_parameter = "Flow"
    elif parameter == "Stage":
        y_label_txt = "Water Surface Elevation (ft)"
        sim_parameter = "Water Surface"
        obs_parameter = "Stage"

    # Units of degrees for EPSG:4326 to buffer outwards from the reference line location
    buffer_increment = 0.001  # Ex: 0.001 degrees is approximately 100 meters
    # Loop through each reference line within the model
    print("Plotting hydrographs for each reference line")
    images_dict = {}
    metrics_list = []
    df_gages_usgs = df_gages_usgs.reset_index(drop=True)
    for idx, line_id in enumerate(ref_lines.refln_id.values):
        # Intersect the gages with the reference lines
        gage_df = df_gages_usgs[
            df_gages_usgs.within(ref_lines.geometry.buffer(buffer_increment).iloc[idx])
        ]
        # Trivial scenario: one row returned indicating one single gage is within the buffer distance
        if len(gage_df) == 1:
            # Site metadata
            usgs_site_name = gage_df.station_nm.values[0]
            usgs_site_id = gage_df.site_no.values[0]
            station_id = f"USGS-{usgs_site_id}"
            usgs_datum = gage_df.alt_va.values[0]
            gage_idx = gage_df.index[0] + 1
            print(f"Plotting {station_id} for reference line {line_id}")

            # Modeled streamflow
            qsim_df = (
                ref_lines_ds.sel(refln_id=line_id)
                [sim_parameter].to_dataframe()[sim_parameter]
                .to_frame()
            )
            qsim_df.columns = ["Modeled"]

            # Observed streamflow: instantaneous values
            qobs_df = get_nwis(usgs_site_id, obs_parameter, 'iv', path_start_date, path_end_date)
            if qobs_df is None:
                print(
                    f"Instantaneous data for USGS station {usgs_site_id} is not available for the calibration period"
                )
                # Observed streamflow: daily values
                qobs_df = get_nwis(usgs_site_id, obs_parameter, 'dv', path_start_date, path_end_date)
                if qobs_df is None:
                    print(
                        f"Daily data for USGS station {usgs_site_id} is not available for the calibration period"
                    )
                    continue

            if parameter == "Stage":
                qobs_df = qobs_df + usgs_datum

            qobs_df.columns = ["Observed"]

            # Resample the data to the same timestep frequency of the model
            qobs_df, qsim_df, timestep = format_datetime(qobs_df, qsim_df)

            # Calculate the metrics between the observed and modeled streamflow
            q_df = pd.merge(
                qobs_df, qsim_df, left_index=True, right_index=True
            ).dropna()
            qobs_df, qsim_df = None, None
            metrics = calc_metrics(q_df, usgs_site_id)

            # Generate the figure
            fig, ax = plt.subplots(figsize=(10, 10))
            # Plot the modeled vs observed streamflow
            q_df["Observed"].plot(ax=ax, color="blue", label="Observed", alpha=0.7)
            q_df["Modeled"].plot(ax=ax, color="red", label="Modeled", alpha=0.7)
            # Add grid lines
            ax.grid()
            # Add a legend
            ax.legend()
            # Add axis labels
            ax.set_xlabel("Date")
            ax.set_ylabel(y_label_txt)
            # Add a title
            ax.set_title(f"USGS-{usgs_site_id} {usgs_site_name}")
            # Set the custom y-axis formatter
            ax.get_yaxis().set_major_formatter(
                ticker.FuncFormatter(lambda x, p: format(int(x), ","))
            )
            # Save the figure
            path = f"{domain_name}_{usgs_site_id}_plan0{plan_index}_{parameter}.png"
            image_path = os.path.join(root_dir, path)
            fig.savefig(image_path, bbox_inches="tight")
            plt.close(fig)
            if (
                report_document is None
                and report_keywords is None
            ):
                images_dict[station_id] = image_path
                metrics_list.append(metrics)
            else:
                # Search for the keyword within the document and add the image above it
                report_document = add_image_to_keyword(
                    report_document,
                    f"«plan0{plan_index}_figure_gage_{parameter}»",
                    image_path,
                )
                os.remove(image_path)
                metrics_list.append(metrics)
                # Update the report text for Table 9: Two-Dimensional Computational Solver Tolerances and Settings
                report_keywords = fill_computation_settings_table(report_keywords, plan_params, plan_attrs, plan_index)
                # Update the report text for Table 11: Gage Calibration Timesteps
                report_keywords[f'table11_gage0{gage_idx}_name'] = usgs_site_name
                report_keywords[f'plan0{plan_index}_gage0{gage_idx}_{parameter.lower()}_ts'] = timestep
                
        # Non-trivial scenario: no rows returned indicating no gages are within the buffer distance
        # Need to buffer out the reference line until the closest gage (within reason) is found
        elif len(gage_df) == 0:
            # Increment the buffer distance
            while len(gage_df) == 0:
                buffer_increment += 0.001
                if buffer_increment > 0.01:
                    print(
                        f"No gages found within ~1-km of the reference line: {line_id}."
                    )
                    break
                else:
                    print(
                        f"No gages found near the reference line {line_id}. Incrementing buffer distance by ~100-m."
                    )
                    gage_df = df_gages_usgs[
                        df_gages_usgs.within(
                            ref_lines.geometry.buffer(buffer_increment).iloc[idx]
                        )
                    ]
            if len(gage_df) == 1:
                # Site metadata
                usgs_site_name = gage_df.station_nm.values[0]
                usgs_site_id = gage_df.site_no.values[0]
                station_id = f"USGS-{usgs_site_id}"
                usgs_datum = gage_df.alt_va.values[0]
                gage_idx = gage_df.index[0] + 1
                print(f"Gage index: {gage_idx}")
                print(f"Plotting {station_id} for reference line {line_id}")

                # Modeled streamflow
                qsim_df = (
                    ref_lines_ds.sel(refln_id=line_id)
                    [sim_parameter].to_dataframe()[sim_parameter]
                    .to_frame()
                )
                qsim_df.columns = ["Modeled"]

                # Observed streamflow
                qobs_df = get_nwis(usgs_site_id, obs_parameter, path_start_date, path_end_date)
                if qobs_df is None:
                    print(
                        f"USGS station {usgs_site_id} is not available for the calibration period"
                    )
                    continue

                if parameter == "Stage":
                    qobs_df = qobs_df + usgs_datum

                qobs_df.columns = ["Observed"]

                # Resample the data to the same timestep frequency of the model
                qobs_df, qsim_df, timestep = format_datetime(qobs_df, qsim_df)

                # Calculate the metrics between the observed and modeled streamflow
                q_df = pd.merge(
                    qobs_df, qsim_df, left_index=True, right_index=True
                ).dropna()
                qobs_df, qsim_df = None, None
                metrics = calc_metrics(q_df, usgs_site_id)

                # Generate the figure
                fig, ax = plt.subplots(figsize=(10, 10))
                # Plot the modeled vs observed streamflow
                q_df["Observed"].plot(
                    ax=ax, color="blue", label="Observed", alpha=0.7
                )
                q_df["Modeled"].plot(ax=ax, color="red", label="Modeled", alpha=0.7)
                # Add grid lines
                ax.grid()
                # Add a legend
                ax.legend()
                # Add axis labels
                ax.set_xlabel("Date")
                ax.set_ylabel(y_label_txt)
                # Add a title
                ax.set_title(f"USGS-{usgs_site_id} {usgs_site_name}")
                # Set the custom y-axis formatter
                ax.get_yaxis().set_major_formatter(
                    ticker.FuncFormatter(lambda x, p: format(int(x), ","))
                )
                # Save the figure
                path = f"{domain_name}_{usgs_site_id}_plan0{plan_index}_{parameter}.png"
                image_path = os.path.join(root_dir,path)
                fig.savefig(image_path, bbox_inches="tight")
                plt.close(fig)
                if (
                    report_document is None
                    and report_keywords is None
                ):
                    images_dict[station_id] = image_path
                    metrics_list.append(metrics)
                else:
                    # Search for the keyword within the document and add the image above it
                    report_document = add_image_to_keyword(
                        report_document,
                        f"«plan0{plan_index}_figure_gage_{parameter}»",
                        image_path,
                    )
                    os.remove(image_path)
                    metrics_list.append(metrics)
                    # Update the report text for Table 9: Two-Dimensional Computational Solver Tolerances and Settings
                    report_keywords = fill_computation_settings_table(report_keywords, plan_params, plan_attrs, plan_index)
                    # Update the report text for Table 11: Gage Calibration Timesteps
                    report_keywords[f'table11_gage0{gage_idx}_name'] = usgs_site_name
                    report_keywords[f'plan0{plan_index}_gage0{gage_idx}_{parameter.lower()}_ts'] = timestep
            elif len(gage_df) > 1:
                raise ValueError(
                    f"{len(gage_df)} gages found within the minimum buffer distance of the reference line. {gage_df}."
                )
    if report_document is None and report_keywords is None:
        # Combine the metrics into a single dataframe
        if len(metrics_list) == 0:
            return images_dict, None
        else:
            metrics_df = pd.concat(metrics_list)
            return images_dict, metrics_df
    else:
        # Update the report text
        report_keywords[
            f"plan0{plan_index}_figure_gage_{parameter}"
        ] = "USGS Gage Flow Calibration Plots"
        report_keywords[
            f"plan0{plan_index}_date"
        ] = f"{path_start_date} to {path_end_date}"
        report_keywords
        if len(metrics_list) == 0:
            return report_document, report_keywords
        else:
            # Combine all gage metrics into a single dataframe
            metrics_df = pd.concat(metrics_list)
            # Update the calibration metrics table in the report
            report_keywords = fill_calibration_metrics_table(
                report_keywords, plan_index, metrics_df, parameter
            )
            return report_document, report_keywords
